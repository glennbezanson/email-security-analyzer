"""
Export Dialog
Export reports in various formats
"""

from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QGroupBox,
    QRadioButton, QCheckBox, QLabel, QPushButton,
    QFileDialog, QProgressBar, QMessageBox
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from pathlib import Path
import json
from datetime import datetime

from ..styles import EdgeColors


class ExportWorker(QThread):
    """Worker thread for export operations"""

    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    progress = pyqtSignal(int)

    def __init__(self, data, format_type, filepath, options):
        super().__init__()
        self.data = data
        self.format_type = format_type
        self.filepath = filepath
        self.options = options

    def run(self):
        try:
            if self.format_type == "json":
                self._export_json()
            elif self.format_type == "csv":
                self._export_csv()
            elif self.format_type == "excel":
                self._export_excel()
            elif self.format_type == "docx":
                self._export_docx()

            self.finished.emit(str(self.filepath))
        except Exception as e:
            self.error.emit(str(e))

    def _export_json(self):
        """Export as JSON"""
        export_data = {}

        if self.options.get('threats'):
            export_data['threats'] = [
                t.to_dict() if hasattr(t, 'to_dict') else str(t)
                for t in self.data.get('threats', [])
            ]

        if self.options.get('cases'):
            export_data['cases'] = [
                c.to_dict() if hasattr(c, 'to_dict') else str(c)
                for c in self.data.get('cases', [])
            ]

        if self.options.get('diagnostics'):
            export_data['diagnostics'] = [
                d.to_dict() if hasattr(d, 'to_dict') else str(d)
                for d in self.data.get('diagnostics', [])
            ]

        export_data['exported_at'] = datetime.utcnow().isoformat()

        with open(self.filepath, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, default=str)

    def _export_csv(self):
        """Export as CSV"""
        import csv

        threats = self.data.get('threats', [])
        if not threats:
            raise ValueError("No threats to export")

        with open(self.filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)

            # Header
            writer.writerow([
                'Threat ID', 'Subject', 'From', 'To', 'Received',
                'Attack Type', 'Severity', 'Status'
            ])

            # Data
            for threat in threats:
                writer.writerow([
                    threat.threat_id,
                    threat.subject,
                    threat.from_address,
                    ', '.join(threat.to_addresses[:3]),
                    threat.received_time.isoformat(),
                    threat.attack_type.value,
                    threat.severity.value,
                    threat.remediation_status.value
                ])

    def _export_excel(self):
        """Export as Excel"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
        except ImportError:
            raise ImportError("openpyxl is required for Excel export")

        wb = Workbook()

        # Threats sheet
        if self.options.get('threats'):
            ws = wb.active
            ws.title = "Threats"

            # Header
            headers = ['Threat ID', 'Subject', 'From', 'To', 'Received',
                      'Attack Type', 'Severity', 'Status']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="486D87", fill_type="solid")

            # Data
            for row, threat in enumerate(self.data.get('threats', []), 2):
                ws.cell(row=row, column=1, value=threat.threat_id)
                ws.cell(row=row, column=2, value=threat.subject)
                ws.cell(row=row, column=3, value=threat.from_address)
                ws.cell(row=row, column=4, value=', '.join(threat.to_addresses[:3]))
                ws.cell(row=row, column=5, value=threat.received_time.isoformat())
                ws.cell(row=row, column=6, value=threat.attack_type.value)
                ws.cell(row=row, column=7, value=threat.severity.value)
                ws.cell(row=row, column=8, value=threat.remediation_status.value)

        # Diagnostics sheet
        if self.options.get('diagnostics') and self.data.get('diagnostics'):
            ws = wb.create_sheet("Diagnostics")

            headers = ['Rule', 'Category', 'Severity', 'Status', 'Evidence']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)

            for row, diag in enumerate(self.data.get('diagnostics', []), 2):
                ws.cell(row=row, column=1, value=diag.rule_name)
                ws.cell(row=row, column=2, value=diag.category)
                ws.cell(row=row, column=3, value=diag.severity)
                ws.cell(row=row, column=4, value="PASS" if diag.passed else "FAIL")
                ws.cell(row=row, column=5, value=diag.evidence)

        wb.save(self.filepath)

    def _export_docx(self):
        """Export as Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            raise ImportError("python-docx is required for Word export")

        doc = Document()

        # Title
        title = doc.add_heading('Email Security Report', 0)

        # Date
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # Threats section
        if self.options.get('threats'):
            doc.add_heading('Threats', level=1)

            threats = self.data.get('threats', [])
            doc.add_paragraph(f"Total threats: {len(threats)}")

            # Summary table
            if threats:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'

                header_cells = table.rows[0].cells
                header_cells[0].text = 'Severity'
                header_cells[1].text = 'Attack Type'
                header_cells[2].text = 'Subject'
                header_cells[3].text = 'Status'

                for threat in threats[:20]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = threat.severity.value
                    row_cells[1].text = threat.attack_type.value
                    row_cells[2].text = threat.subject[:50]
                    row_cells[3].text = threat.remediation_status.value

        # Diagnostics section
        if self.options.get('diagnostics'):
            doc.add_heading('Diagnostics', level=1)

            diagnostics = self.data.get('diagnostics', [])
            passed = sum(1 for d in diagnostics if d.passed)
            failed = len(diagnostics) - passed

            doc.add_paragraph(f"Passed: {passed}, Failed: {failed}")

            # Failed items
            for diag in diagnostics:
                if not diag.passed:
                    doc.add_heading(diag.rule_name, level=2)
                    doc.add_paragraph(f"Severity: {diag.severity}")
                    doc.add_paragraph(f"Evidence: {diag.evidence}")

        doc.save(self.filepath)


class ExportDialog(QDialog):
    """
    Dialog for exporting reports
    """

    def __init__(self, parent=None):
        super().__init__(parent)

        self.main_window = parent
        self.setWindowTitle("Export Report")
        self.setMinimumSize(400, 350)

        self._setup_ui()

    def _setup_ui(self):
        """Setup dialog UI"""
        layout = QVBoxLayout(self)
        layout.setSpacing(16)

        # Format selection
        format_group = QGroupBox("Export Format")
        format_layout = QVBoxLayout()

        self.json_radio = QRadioButton("JSON (machine-readable)")
        self.json_radio.setChecked(True)
        format_layout.addWidget(self.json_radio)

        self.csv_radio = QRadioButton("CSV (spreadsheet)")
        format_layout.addWidget(self.csv_radio)

        self.excel_radio = QRadioButton("Excel (.xlsx)")
        format_layout.addWidget(self.excel_radio)

        self.docx_radio = QRadioButton("Word Document (.docx)")
        format_layout.addWidget(self.docx_radio)

        format_group.setLayout(format_layout)
        layout.addWidget(format_group)

        # Content selection
        content_group = QGroupBox("Include")
        content_layout = QVBoxLayout()

        self.threats_check = QCheckBox("Threats")
        self.threats_check.setChecked(True)
        content_layout.addWidget(self.threats_check)

        self.cases_check = QCheckBox("Cases")
        self.cases_check.setChecked(True)
        content_layout.addWidget(self.cases_check)

        self.diagnostics_check = QCheckBox("Diagnostic Results")
        self.diagnostics_check.setChecked(True)
        content_layout.addWidget(self.diagnostics_check)

        content_group.setLayout(content_layout)
        layout.addWidget(content_group)

        # Progress
        self.progress = QProgressBar()
        self.progress.hide()
        layout.addWidget(self.progress)

        # Buttons
        buttons = QHBoxLayout()
        buttons.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setObjectName("secondary")
        cancel_btn.clicked.connect(self.reject)
        buttons.addWidget(cancel_btn)

        self.export_btn = QPushButton("Export...")
        self.export_btn.clicked.connect(self._do_export)
        buttons.addWidget(self.export_btn)

        layout.addLayout(buttons)

    def _do_export(self):
        """Perform export"""
        # Get format
        if self.json_radio.isChecked():
            format_type = "json"
            filter_str = "JSON Files (*.json)"
            extension = ".json"
        elif self.csv_radio.isChecked():
            format_type = "csv"
            filter_str = "CSV Files (*.csv)"
            extension = ".csv"
        elif self.excel_radio.isChecked():
            format_type = "excel"
            filter_str = "Excel Files (*.xlsx)"
            extension = ".xlsx"
        else:
            format_type = "docx"
            filter_str = "Word Documents (*.docx)"
            extension = ".docx"

        # Get filename
        filename = f"security_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}{extension}"
        filepath, _ = QFileDialog.getSaveFileName(
            self,
            "Export Report",
            filename,
            filter_str
        )

        if not filepath:
            return

        # Gather data
        data = {}
        if self.threats_check.isChecked():
            data['threats'] = getattr(self.main_window.threats_view, 'threats', [])
        if self.cases_check.isChecked():
            data['cases'] = getattr(self.main_window.cases_view, 'cases', [])
        if self.diagnostics_check.isChecked():
            data['diagnostics'] = getattr(self.main_window.diagnostics_view, 'results', [])

        options = {
            'threats': self.threats_check.isChecked(),
            'cases': self.cases_check.isChecked(),
            'diagnostics': self.diagnostics_check.isChecked()
        }

        # Start export
        self.progress.setMaximum(0)
        self.progress.show()
        self.export_btn.setEnabled(False)

        self.worker = ExportWorker(data, format_type, filepath, options)
        self.worker.finished.connect(self._on_export_finished)
        self.worker.error.connect(self._on_export_error)
        self.worker.start()

    def _on_export_finished(self, filepath: str):
        """Handle export completion"""
        self.progress.hide()
        self.export_btn.setEnabled(True)

        QMessageBox.information(
            self,
            "Export Complete",
            f"Report exported successfully to:\n{filepath}"
        )
        self.accept()

    def _on_export_error(self, error: str):
        """Handle export error"""
        self.progress.hide()
        self.export_btn.setEnabled(True)

        QMessageBox.critical(
            self,
            "Export Failed",
            f"Failed to export report:\n{error}"
        )
